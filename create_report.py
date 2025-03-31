from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('BÁO CÁO BÀI TẬP LỚN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin bài tập
    doc.add_heading('Đề tài: Trợ giúp công ty tuyển dụng ứng viên xin việc', 1)
    
    # Mục lục
    doc.add_heading('Mục lục', 1)
    toc = [
        "1. Thiết kế cơ sở dữ liệu",
        "2. Phát biểu bài toán",
        "3. Mô hình toán học",
        "4. Phương pháp",
        "5. Cài đặt"
    ]
    for item in toc:
        doc.add_paragraph(item)
    
    # 1. Thiết kế cơ sở dữ liệu
    doc.add_heading('1. Thiết kế cơ sở dữ liệu', 1)
    doc.add_paragraph('''
    Hệ thống được thiết kế với các bảng dữ liệu chính sau:
    
    - Bảng Jobs (Công việc):
        + id: Khóa chính
        + title: Tiêu đề công việc
        + description: Mô tả công việc
        + required_skills: Yêu cầu kỹ năng
        + required_experience: Kinh nghiệm yêu cầu
        + offered_salary: Mức lương
        + max_candidates: Số lượng ứng viên tối đa
        + industry: Ngành nghề
    
    - Bảng Candidates (Ứng viên):
        + id: Khóa chính
        + name: Họ tên
        + skills: Kỹ năng
        + experience: Kinh nghiệm
        + desired_salary: Mức lương mong muốn
        + industry: Ngành nghề
    ''')
    
    # 2. Phát biểu bài toán
    doc.add_heading('2. Phát biểu bài toán', 1)
    doc.add_paragraph('''
    Bài toán: Hệ thống cần giải quyết vấn đề quản lý và phân bổ ứng viên cho các công việc phù hợp dựa trên thông tin của họ.
    
    Đầu vào:
    - Thông tin ứng viên (kỹ năng, kinh nghiệm)
    - Thông tin công việc (yêu cầu kỹ năng, mức lương)
    - Đơn ứng tuyển
    
    Đầu ra:
    - Danh sách ứng viên phù hợp cho từng công việc
    - Trạng thái đơn ứng tuyển
    
    Biến môi trường:
    - Số lượng công việc tối đa
    - Số lượng ứng viên tối đa
    
    Biến quyết định:
    - Quyết định phân bổ ứng viên vào công việc nào
    
    Ràng buộc:
    - Ứng viên chỉ được làm một công việc
    - Công việc có giới hạn số lượng ứng viên
    
    Hàm mục tiêu:
    - Tối đa hóa mức độ phù hợp giữa ứng viên và công việc
    ''')
    
    # 3. Mô hình toán học
    doc.add_heading('3. Mô hình toán học', 1)
    doc.add_paragraph('''
    Ký hiệu:
    - Tập hợp: Ứng viên (C), công việc (J), kỹ năng (S)
    - Biến quyết định: x_ij (1 nếu ứng viên i được chọn cho công việc j, 0 nếu không)
    - Tham số: Trình độ kỹ năng, kinh nghiệm, mức lương
    
    Hàm mục tiêu:
    max Σ (phù hợp_ij * x_ij)
    
    Ràng buộc:
    1. Mỗi ứng viên chỉ làm một công việc: Σ x_ij ≤ 1 với mọi i
    2. Số lượng ứng viên cho mỗi công việc có giới hạn
    3. Yêu cầu về kỹ năng, kinh nghiệm, ngành nghề, địa điểm, mức lương
    ''')
    
    # 4. Phương pháp
    doc.add_heading('4. Phương pháp', 1)
    doc.add_paragraph('''
    Hệ thống sử dụng thuật toán TOPSIS (Technique for Order Preference by Similarity to an Ideal Solution) để đánh giá và phân bổ ứng viên:
    
    Bước 1: Chuẩn hóa dữ liệu
    - Đưa các giá trị về thang đo [0,1]
    - Xử lý các thuộc tính định tính và định lượng
    
    Bước 2: Tính ma trận trọng số
    - Gán trọng số cho các tiêu chí
    - Tính ma trận trọng số v = r * w
    
    Bước 3: Xác định giải pháp lý tưởng
    - Giải pháp lý tưởng dương A*
    - Giải pháp lý tưởng âm A-
    
    Bước 4: Tính khoảng cách
    - Khoảng cách đến giải pháp lý tưởng dương S*
    - Khoảng cách đến giải pháp lý tưởng âm S-
    
    Bước 5: Tính độ tương tự
    - C* = S- / (S* + S-)
    
    Bước 6: Phân bổ ứng viên
    - Sắp xếp theo độ tương tự giảm dần
    - Phân bổ theo ràng buộc
    ''')
    
    # 5. Cài đặt
    doc.add_heading('5. Cài đặt', 1)
    doc.add_paragraph('''
    Công nghệ sử dụng:
    - Backend: Python với Flask
    - Database: SQLite
    - Thuật toán: TOPSIS cho việc đánh giá và phân bổ
    
    Cấu trúc dự án:
    recruitment_system/
    ├── models/
    │   ├── database.py: Xử lý kết nối database
    │   └── allocation.py: Triển khai thuật toán TOPSIS
    └── requirements.txt: Các thư viện cần thiết
    
    Các chức năng chính:
    1. Thêm Job
    2. Thêm Candidates
    3. Hiển thị thông tin chi tiết qua modal
    4. Đánh giá và phân bổ ứng viên tự động
    ''')
    
    # Lưu file
    doc.save('Bao_Cao_Bai_Tap_Lon.docx')

if __name__ == '__main__':
    create_report() 